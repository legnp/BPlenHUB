import openpyxl
import docx
import json
import os
import shutil
import re
from datetime import datetime

def extract_cell_markdown(cell):
    paragraphs_md = []
    for p in cell.paragraphs:
        p_text = ""
        for r in p.runs:
            text = r.text
            # Detect bold run
            is_bold = r.bold or (r.style and r.style.font and r.style.font.bold)
            if is_bold and text.strip():
                # Avoid wrapping empty spaces in asterisks
                leading_spaces = len(text) - len(text.lstrip())
                trailing_spaces = len(text) - len(text.rstrip())
                mid_text = text.strip()
                p_text += text[:leading_spaces] + f"**{mid_text}**" + (text[len(text)-trailing_spaces:] if trailing_spaces else "")
            else:
                p_text += text
        paragraphs_md.append(p_text)
    return "\n".join(paragraphs_md).strip()

scratch_excel = r"D:\BPlen HUB\v3\scratch\servicos_bplen-v3.xlsx"
portfolio_path = r"D:\BPlen HUB\v3\portfolio\portfolio_bplen.xlsx"
docx_path = r"D:\BPlen HUB\v3\portfolio\anuncios_bplen.docx"
output_payload_path = r"D:\BPlen HUB\v3\portfolio\portfolio_payload.json"

print("=========================================")
print("  BPlen HUB Portfolio Parser & Sync      ")
print("=========================================\n")

# 1. RESTORE EXCEL TO RECOVER FORMULA VALUES CACHE
print("Step 1: Restoring template excel to recover cached formula values...")
if os.path.exists(scratch_excel):
    shutil.copy(scratch_excel, portfolio_path)
    print("Template restored successfully.")
else:
    print(f"WARNING: Template not found at {scratch_excel}. Parsing existing file.")

# Ensure the new sheets exist in the restored file
os.system(f'python "D:\\BPlen HUB\\v3\\scratch\\init_excel_sheets.py"')

# 2. READ COMMERCAL PRICES IN-MEMORY (DATA ONLY)
print("\nStep 2: Extracting commercial pricing from template in data_only mode...")
wb_data = openpyxl.load_workbook(scratch_excel, data_only=True)
wb_config = openpyxl.load_workbook(portfolio_path, data_only=True) # To read Jornada and Checkpoints

# Services Coordinates
services_coords = {
    "BPL-001": {
        "title": "Posicionamento Profissional",
        "slug": "posicionamento-profissional",
        "order": 2,
        "is_journey": True,
        "price_row": 41, # Card price Column E
        "max_inst_row": 43, # Max installments Column D
        "pix_row": 48, # PIX price Column E
        "quotas": {"posicionamentoprofissional": 1}
    },
    "BPL-002": {
        "title": "Análise Comportamental (DISC)",
        "slug": "analise-comportamental",
        "order": 3,
        "is_journey": True,
        "price_row": 72,
        "max_inst_row": 74,
        "pix_row": 79,
        "quotas": {"analisecomportamental": 1, "devolutiva-analise-comportamental": 1}
    },
    "BPL-003": {
        "title": "Plano de Carreira",
        "slug": "plano-de-carreira",
        "order": 4,
        "is_journey": True,
        "price_row": 102,
        "max_inst_row": 104,
        "pix_row": 109,
        "quotas": {"planodecarreira": 1, "consultoria-plano-carreira": 1}
    },
    "BPL-004": {
        "title": "Gestão e Desenvolvimento (GDC)",
        "slug": "gestao-e-desenvolvimento",
        "order": 5,
        "is_journey": True,
        "price_row": 133,
        "max_inst_row": 135,
        "pix_row": 140,
        "quotas": {"gestaoedesenvolvimento": 1, "1-to-1": 10}
    },
    "BPL-005": {
        "title": "MentoCoach (Alta Performance)",
        "slug": "mentocoach",
        "order": 6,
        "is_journey": True,
        "price_row": 162,
        "max_inst_row": 164,
        "pix_row": 169,
        "quotas": {"mentocoach": 1, "1-to-1": 12}
    }
}

custos_sheet = wb_data["Custo de Serviços"]
services_data = {}

for code, cfg in services_coords.items():
    price_val = custos_sheet.cell(row=cfg["price_row"], column=5).value
    max_inst_val = custos_sheet.cell(row=cfg["max_inst_row"], column=4).value
    pix_val = custos_sheet.cell(row=cfg["pix_row"], column=5).value
    
    price = round(float(price_val), 2) if price_val is not None else 0.0
    max_installments = int(max_inst_val) if max_inst_val is not None else 12
    price_pix = round(float(pix_val), 2) if pix_val is not None else 0.0
    
    services_data[code] = {
        "id": cfg["slug"],
        "slug": cfg["slug"],
        "serviceCode": code,
        "title": cfg["title"],
        "price": price,
        "pricePix": price_pix,
        "maxInstallments": max_installments,
        "isStepJourney": cfg["is_journey"],
        "order": cfg["order"],
        "grantedQuotas": cfg["quotas"],
        "targetAudiences": ["people", "companies"],
        "status": "active"
    }
    print(f" -> Parsed Service {code}: Card R$ {price:.2f} | PIX R$ {price_pix:.2f} | {max_installments}x")

# Packages Coordinates
pacotes_sheet = wb_data["Pacotes de Serviço"]
packages_data = {}
package_coords = {
    "BPL-PAC-JR": {"name": "Pacote JUNIOR", "slug": "pacote-junior", "row": 8, "quotas": {"posicionamentoprofissional": 1}},
    "BPL-PAC-PL": {"name": "Pacote PLENO", "slug": "pacote-pleno", "row": 9, "quotas": {"posicionamentoprofissional": 1, "analisecomportamental": 1, "devolutiva-analise-comportamental": 1}},
    "BPL-PAC-SR": {"name": "Pacote SENIOR (Recomendado)", "slug": "pacote-senior", "row": 10, "quotas": {"posicionamentoprofissional": 1, "analisecomportamental": 1, "devolutiva-analise-comportamental": 1, "planodecarreira": 1, "consultoria-plano-carreira": 1}},
    "BPL-PAC-LD": {"name": "Pacote LIDER", "slug": "pacote-lider", "row": 11, "quotas": {"posicionamentoprofissional": 1, "analisecomportamental": 1, "devolutiva-analise-comportamental": 1, "planodecarreira": 1, "consultoria-plano-carreira": 1, "gestaoedesenvolvimento": 1, "1-to-1": 10}},
    "BPL-PAC-EB": {"name": "Pacote EMBAIXADOR BPLEN", "slug": "pacote-embaixador", "row": 12, "quotas": {"posicionamentoprofissional": 1, "analisecomportamental": 1, "devolutiva-analise-comportamental": 1, "planodecarreira": 1, "consultoria-plano-carreira": 1, "gestaoedesenvolvimento": 1, "mentocoach": 1, "1-to-1": 22}}
}

for code, cfg in package_coords.items():
    r = cfg["row"]
    price_card = pacotes_sheet.cell(row=r, column=6).value
    price_pix = pacotes_sheet.cell(row=r, column=7).value
    
    price = round(float(price_card), 2) if price_card is not None else 0.0
    price_pix = round(float(price_pix), 2) if price_pix is not None else 0.0
    
    packages_data[code] = {
        "id": cfg["slug"],
        "slug": cfg["slug"],
        "serviceCode": code,
        "title": cfg["name"],
        "price": price,
        "pricePix": price_pix,
        "maxInstallments": 12,
        "isStepJourney": False,
        "grantedQuotas": cfg["quotas"],
        "targetAudiences": ["people", "companies"],
        "status": "active"
    }
    print(f" -> Parsed Package {code}: Card R$ {price:.2f} | PIX R$ {price_pix:.2f}")

wb_data.close()


# 3. READ JOURNEY CONFIGURATION (ORDER)
print("\nStep 3: Reading Journey orders from 'Jornada' sheet...")
try:
    journey_sheet = wb_config["Jornada"]
    journey_orders = {}
    for r in range(2, journey_sheet.max_row + 1):
        code = str(journey_sheet.cell(row=r, column=1).value or "").strip()
        order = journey_sheet.cell(row=r, column=2).value
        if code and order is not None:
            journey_orders[code] = int(order)
    print(f" -> Loaded orders for {len(journey_orders)} services.")
except KeyError:
    print("WARNING: 'Jornada' sheet not found. Using default orders.")
    journey_orders = {
        "BPL-000": 1, "BPL-001": 2, "BPL-002": 3, "BPL-003": 4, 
        "BPL-004": 5, "BPL-005": 6, "BPL-006": 7
    }

# Update services_data with dynamic orders
for code, data in services_data.items():
    if code in journey_orders:
        data["order"] = journey_orders[code]
        print(f"    * {code} order set to: {data['order']}")


# 4. PARSE CHECKPOINTS FROM EXCEL (READ ONLY)
print("\nStep 4: Reading checkpoints from 'Checkpoints' sheet...")
try:
    cp_sheet = wb_config["Checkpoints"]
    checkpoints_by_service = {}

    for r in range(2, cp_sheet.max_row + 1):
        service_code = cp_sheet.cell(row=r, column=1).value
        if not service_code:
            continue
        service_code = str(service_code).strip()
        checkpoint_id = str(cp_sheet.cell(row=r, column=3).value or "").strip()
        order = cp_sheet.cell(row=r, column=4).value
        title = str(cp_sheet.cell(row=r, column=5).value or "").strip()
        type_val = str(cp_sheet.cell(row=r, column=6).value or "").strip()
        ref_id = str(cp_sheet.cell(row=r, column=7).value or "").strip()
        desc = str(cp_sheet.cell(row=r, column=8).value or "").strip()
        
        if service_code not in checkpoints_by_service:
            checkpoints_by_service[service_code] = []
            
        checkpoints_by_service[service_code].append({
            "id": f"ss-{type_val}-{ref_id}",
            "type": type_val,
            "referenceId": ref_id,
            "title": title,
            "description": desc or "Etapa recomendada de desenvolvimento"
        })
    print(f" -> Loaded checkpoints for {len(checkpoints_by_service)} services.")
except KeyError:
    print("ERROR: 'Checkpoints' sheet not found. Delivery steps will be empty.")
    checkpoints_by_service = {}

wb_config.close()


# 5. PARSE WORD ADVERTISEMENTS (COPYWRITING)
print("\nStep 5: Extracting advertisements and FAQs from anuncios_bplen.docx...")
doc = docx.Document(docx_path)

# Services Copywriting (Tables 0-4)
for idx, code in enumerate(["BPL-001", "BPL-002", "BPL-003", "BPL-004", "BPL-005"]):
    table = doc.tables[idx]
    
    kicker = table.rows[1].cells[1].text.strip()
    service_title = table.rows[2].cells[1].text.strip()
    short_desc = extract_cell_markdown(table.rows[3].cells[1])
    long_desc = extract_cell_markdown(table.rows[4].cells[1])
    faq_raw = extract_cell_markdown(table.rows[6].cells[1])
    workflow_raw = extract_cell_markdown(table.rows[7].cells[1])
    
    # Parse FAQ array
    faq_list = []
    faq_matches = re.findall(r"Q\d+:\s*(.*?)\nA\d+:\s*(.*?)(?=\n+Q\d+:|$)", faq_raw, re.DOTALL)
    for q, a in faq_matches:
        faq_list.append({
            "question": q.strip(),
            "answer": a.strip()
        })
        
    # Parse workflow steps
    workflow_list = []
    flow_items = re.findall(r"\d+\.\s*([^;\n]+)", workflow_raw)
    for w_idx, f_title in enumerate(flow_items):
        workflow_list.append({
            "id": f"wf-{code.lower()}-{w_idx + 1}",
            "title": f_title.strip(),
            "type": "task",
            "description": "Etapa recomendada de desenvolvimento"
        })
        
    if code in services_data:
        services_data[code].update({
            "title": service_title,
            "kicker": kicker,
            "sheet": {
                "description": long_desc,
                "coverImage": f"/images/products/{services_data[code]['slug']}.png",
                "paymentConditions": f"Pagamento facilitado no cartão em até {services_data[code]['maxInstallments']}x ou PIX com desconto especial.",
                "faq": faq_list,
                "termsAndConditions": "Ao contratar este serviço, você concorda com o plano de entrega e as diretrizes do BPlen HUB.",
                "seo": {
                    "title": f"{service_title} | BPlen HUB",
                    "description": short_desc,
                    "keywords": [kicker.lower(), "bplen", "carreira", "profissional"]
                }
            },
            "workflow": workflow_list,
            "deliverySteps": checkpoints_by_service.get(code, []),
            "capabilities": {
                "surveys": [cp["referenceId"] for cp in checkpoints_by_service.get(code, []) if cp["type"] == "survey"],
                "forms": [cp["referenceId"] for cp in checkpoints_by_service.get(code, []) if cp["type"] == "form"],
                "allowedEventTypes": [cp["referenceId"] for cp in checkpoints_by_service.get(code, []) if cp["type"] == "meeting"]
            }
        })
        print(f" -> Integrated Copywriting for {code} ({service_title}) - FAQs: {len(faq_list)}, Workflow items: {len(workflow_list)}")

# Packages Copywriting (Tables 5-9)
for idx, code in enumerate(["BPL-PAC-JR", "BPL-PAC-PL", "BPL-PAC-SR", "BPL-PAC-LD", "BPL-PAC-EB"]):
    table = doc.tables[idx + 5]
    
    tatic = extract_cell_markdown(table.rows[0].cells[1])
    slogan = extract_cell_markdown(table.rows[1].cells[1])
    kicker = table.rows[2].cells[1].text.strip()
    
    if code in packages_data:
        packages_data[code].update({
            "kicker": kicker,
            "sheet": {
                "description": f"{tatic}\n\n{slogan}",
                "coverImage": f"/images/products/{packages_data[code]['slug']}.png",
                "paymentConditions": "Parcele em até 12x sem juros no cartão de crédito ou obtenha desconto exclusivo via PIX.",
                "faq": [
                    {"question": "O que está incluso neste pacote?", "answer": f"Este pacote engloba um conjunto estratégico de serviços unificados, eliminando custos redundantes: {tatic}"},
                    {"question": "Como funciona o agendamento?", "answer": "Assim que a contratação for concluída, as sessões e relatórios correspondentes serão desbloqueados na sua Área de Membro para livre agendamento."}
                ],
                "termsAndConditions": "Os pacotes estruturados oferecem licenças de uso combinadas e condições comerciais unificadas.",
                "seo": {
                    "title": f"{packages_data[code]['title']} | BPlen HUB",
                    "description": slogan,
                    "keywords": ["pacote bplen", "aceleração profissional", "carreira executiva"]
                }
            },
            "workflow": [],
            "deliverySteps": [],
            "capabilities": {
                "surveys": [],
                "forms": [],
                "allowedEventTypes": []
            }
        })
        print(f" -> Integrated Copywriting for Package {code} ({kicker})")


# 6. DEFINE INTERNAL/UTILITY SERVICES (BPL-000 AND BPL-006)
print("\nStep 6: Adding internal system services (BPL-000 and BPL-006)...")
internal_services = {
    "BPL-000": {
        "id": "onboarding",
        "slug": "onboarding",
        "serviceCode": "BPL-000",
        "title": "Onboarding Estratégico",
        "kicker": "Onboarding",
        "price": 0.0,
        "pricePix": 0.0,
        "maxInstallments": 12,
        "isStepJourney": True,
        "order": 1,
        "grantedQuotas": {"onboarding": 1},
        "targetAudiences": ["people", "internal"],
        "status": "active",
        "sheet": {
            "description": "Dê os seus primeiros passos no BPlen HUB e conheça todo o nosso ecossistema.",
            "coverImage": "/images/products/onboarding.png",
            "paymentConditions": "Serviço interno gratuito para membros.",
            "faq": [
                {"question": "Como iniciar?", "answer": "Basta assistir ao vídeo de introdução e realizar o seu check-in inicial."}
            ],
            "termsAndConditions": "Serviço exclusivo de Onboarding do BPlen HUB.",
            "seo": {
                "title": "Onboarding Estratégico | BPlen HUB",
                "description": "Boas-vindas ao BPlen HUB.",
                "keywords": ["onboarding", "bplen"]
            }
        },
        "workflow": [],
        "deliverySteps": checkpoints_by_service.get("BPL-000", []),
        "capabilities": {
            "surveys": ["check_in"],
            "forms": [],
            "allowedEventTypes": ["onboarding"]
        }
    },
    "BPL-006": {
        "id": "offboarding",
        "slug": "offboarding",
        "serviceCode": "BPL-006",
        "title": "Offboarding e Consolidação",
        "kicker": "Encerramento",
        "price": 0.0,
        "pricePix": 0.0,
        "maxInstallments": 12,
        "isStepJourney": True,
        "order": 7,
        "grantedQuotas": {"offboarding": 1},
        "targetAudiences": ["people", "internal"],
        "status": "active",
        "sheet": {
            "description": "Consolide seus aprendizados, metas e decole para os seus próximos desafios.",
            "coverImage": "/images/products/offboarding.png",
            "paymentConditions": "Serviço interno gratuito para membros em final de ciclo.",
            "faq": [
                {"question": "Qual objetivo?", "answer": "Avaliar o impacto da jornada BPlen e mapear os seus próximos passos autonômos de alta performance."}
            ],
            "termsAndConditions": "Serviço exclusivo de encerramento do BPlen HUB.",
            "seo": {
                "title": "Offboarding e Consolidação | BPlen HUB",
                "description": "Encerramento e consolidação.",
                "keywords": ["offboarding", "bplen"]
            }
        },
        "workflow": [],
        "deliverySteps": checkpoints_by_service.get("BPL-006", []),
        "capabilities": {
            "surveys": ["offboarding_survey"],
            "forms": [],
            "allowedEventTypes": []
        }
    }
}


# 6.B PARSE CAMPAIGNS AND COUPONS (campanhas_bplen.xlsx)
print("\nStep 6.B: Parsing campaigns and coupons from campanhas_bplen.xlsx...")
campanhas_path = r"D:\BPlen HUB\v3\portfolio\campanhas_bplen.xlsx"
campanhas_payload_path = r"D:\BPlen HUB\v3\portfolio\campanhas_payload.json"

code_to_slug = {
    "BPL-000": "onboarding",
    "BPL-001": "posicionamento-profissional",
    "BPL-002": "analise-comportamental",
    "BPL-003": "plano-carreira",
    "BPL-004": "gestao-e-desenvolvimento",
    "BPL-005": "mentocoach",
    "BPL-006": "offboarding",
    "BPL-PAC-JR": "pacote-junior",
    "BPL-PAC-PL": "pacote-pleno",
    "BPL-PAC-SR": "pacote-senior",
    "BPL-PAC-LD": "pacote-lider",
    "BPL-PAC-EB": "pacote-embaixador"
}

if os.path.exists(campanhas_path):
    wb_cam = openpyxl.load_workbook(campanhas_path, data_only=True)
    
    # 1. Process active promotional offers
    if "Ofertas" in wb_cam.sheetnames:
        print(" -> Processing active promotional offers...")
        ofertas_sheet = wb_cam["Ofertas"]
        now_str = datetime.now().strftime("%Y-%m-%d")
        
        for r in range(2, ofertas_sheet.max_row + 1):
            service_code = ofertas_sheet.cell(row=r, column=1).value
            if not service_code:
                continue
            service_code = str(service_code).strip()
            
            promo_pix_val = ofertas_sheet.cell(row=r, column=2).value
            promo_cartao_val = ofertas_sheet.cell(row=r, column=3).value
            data_inicio = ofertas_sheet.cell(row=r, column=4).value
            data_fim = ofertas_sheet.cell(row=r, column=5).value
            slogan_oferta = ofertas_sheet.cell(row=r, column=6).value
            
            # Format dates to string
            start_str = str(data_inicio).split()[0] if data_inicio else ""
            end_str = str(data_fim).split()[0] if data_fim else ""
            
            # Check if campaign is active
            is_active = True
            if start_str and now_str < start_str:
                is_active = False
            if end_str and now_str > end_str:
                is_active = False
                
            if is_active:
                promo_pix = round(float(promo_pix_val), 2) if promo_pix_val is not None else 0.0
                promo_cartao = round(float(promo_cartao_val), 2) if promo_cartao_val is not None else 0.0
                slogan = str(slogan_oferta).strip() if slogan_oferta else ""
                
                # Apply promotion
                product_target = None
                if service_code in services_data:
                    product_target = services_data[service_code]
                elif service_code in packages_data:
                    product_target = packages_data[service_code]
                    
                if product_target:
                    product_target["originalPrice"] = product_target["price"]
                    product_target["originalPricePix"] = product_target["pricePix"]
                    product_target["price"] = promo_cartao
                    product_target["pricePix"] = promo_pix
                    if slogan:
                        product_target["promoLabel"] = slogan
                    print(f"    * Promotion Applied to {service_code}: Card R$ {promo_cartao:.2f} | PIX R$ {promo_pix:.2f} ({slogan})")
                    
    # 2. Process coupons
    coupons_list = []
    if "Cupons" in wb_cam.sheetnames:
        print(" -> Processing coupons...")
        cupons_sheet = wb_cam["Cupons"]
        
        for r in range(2, cupons_sheet.max_row + 1):
            cupom_code = cupons_sheet.cell(row=r, column=1).value
            if not cupom_code:
                continue
            cupom_code = str(cupom_code).strip()
            
            tipo = str(cupons_sheet.cell(row=r, column=2).value or "").strip()
            valor_val = cupons_sheet.cell(row=r, column=3).value
            status = str(cupons_sheet.cell(row=r, column=4).value or "").strip()
            validade = cupons_sheet.cell(row=r, column=5).value
            servicos_hab = str(cupons_sheet.cell(row=r, column=6).value or "").strip()
            
            # Map type and status
            discount_type = "percentage" if tipo == "percentual" else "fixed"
            active_bool = True if status == "ativo" else False
            valor = float(valor_val) if valor_val is not None else 0.0
            
            validade_str = str(validade).split()[0] if validade else ""
            
            restricted_products = []
            if servicos_hab and servicos_hab.lower() != "todos" and servicos_hab.lower() != "all":
                for part in servicos_hab.split(","):
                    clean_part = part.strip()
                    if clean_part in code_to_slug:
                        restricted_products.append(code_to_slug[clean_part])
                        
            coupon_item = {
                "id": cupom_code.upper(),
                "code": cupom_code.upper(),
                "type": discount_type,
                "value": valor,
                "description": f"Desconto de {valor}%" if discount_type == "percentage" else f"Desconto de R$ {valor}",
                "active": active_bool,
                "usageCount": 0
            }
            
            if validade_str:
                coupon_item["expiryDate"] = validade_str
            if restricted_products:
                coupon_item["restrictedToProducts"] = restricted_products
                
            coupons_list.append(coupon_item)
            print(f"    * Coupon Parsed: {coupon_item['code']} | Type: {discount_type} | Value: {valor} | Active: {active_bool}")
            
    # Write coupons payload
    with open(campanhas_payload_path, "w", encoding="utf-8") as f_cam:
        json.dump(coupons_list, f_cam, ensure_ascii=False, indent=2)
    print(f" -> Compiled {len(coupons_list)} coupons to {campanhas_payload_path}")
    wb_cam.close()
else:
    print(f"WARNING: campanhas_bplen.xlsx not found at {campanhas_path}. No campaigns processed.")


# 7. EXPORT COMPILED CATALOGUE TO PAYLOAD JSON
print("\nStep 7: Compiling final payload...")
all_products = []
all_products.append(internal_services["BPL-000"])

for code in ["BPL-001", "BPL-002", "BPL-003", "BPL-004", "BPL-005"]:
    all_products.append(services_data[code])

all_products.append(internal_services["BPL-006"])

for code in ["BPL-PAC-JR", "BPL-PAC-PL", "BPL-PAC-SR", "BPL-PAC-LD", "BPL-PAC-EB"]:
    all_products.append(packages_data[code])

# Ensure target folder exists
os.makedirs(os.path.dirname(output_payload_path), exist_ok=True)

with open(output_payload_path, "w", encoding="utf-8") as f:
    json.dump(all_products, f, ensure_ascii=False, indent=2)

print("\n=========================================")
print("  PARSER PIPELINE EXECUTED WITH SUCCESS! ")
print("=========================================")
print(f"Payload Saved: {output_payload_path}")
print(f"Products Compiled: {len(all_products)} (7 Services | 5 Packages)")
print(f"Protected directory '/portfolio' holds secure sources and outputs.")
print("=========================================")
